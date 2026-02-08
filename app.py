
import gradio as gr
import tempfile
import os
import traceback
import time
import json
import zipfile
import hashlib
import base64
from pathlib import Path
from datetime import datetime
from PIL import Image
import io

import pymupdf4llm
import fitz  # PyMuPDF for metadata and preview
import structure_engine
import markdown2
from docx import Document

# Constants
MIN_TEXT_THRESHOLD = 50
HISTORY_FILE = "conversion_history.json"
CACHE_DIR = "cache"
IMAGES_DIR = "extracted_images"
STATS_FILE = "usage_stats.json"
MAX_HISTORY = 10

# Create directories
os.makedirs(CACHE_DIR, exist_ok=True)
os.makedirs(IMAGES_DIR, exist_ok=True)

def load_history():
    if os.path.exists(HISTORY_FILE):
        try:
            with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return []
    return []

def save_history(entry):
    history = load_history()
    history.insert(0, entry)
    history = history[:MAX_HISTORY]
    with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
        json.dump(history, f, indent=2)

def load_stats():
    if os.path.exists(STATS_FILE):
        try:
            with open(STATS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            pass
    return {"total_conversions": 0, "total_words": 0, "total_files": 0, "avg_time": 0}

def save_stats(stats):
    with open(STATS_FILE, 'w', encoding='utf-8') as f:
        json.dump(stats, f, indent=2)

def update_stats(words, elapsed_time):
    stats = load_stats()
    stats["total_conversions"] += 1
    stats["total_words"] += words
    stats["total_files"] += 1
    n = stats["total_conversions"]
    stats["avg_time"] = ((stats["avg_time"] * (n-1)) + elapsed_time) / n
    save_stats(stats)
    return stats

def get_file_hash(file_path):
    hasher = hashlib.md5()
    with open(file_path, 'rb') as f:
        buf = f.read(65536)
        while buf:
            hasher.update(buf)
            buf = f.read(65536)
    return hasher.hexdigest()

def get_cached_result(file_hash, dpi, export_format):
    cache_key = f"{file_hash}_{dpi}_{export_format}"
    cache_path = os.path.join(CACHE_DIR, f"{cache_key}.json")
    if os.path.exists(cache_path):
        try:
            with open(cache_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            pass
    return None

def save_to_cache(file_hash, dpi, export_format, markdown_text, method_used):
    cache_key = f"{file_hash}_{dpi}_{export_format}"
    cache_path = os.path.join(CACHE_DIR, f"{cache_key}.json")
    with open(cache_path, 'w', encoding='utf-8') as f:
        json.dump({
            "markdown": markdown_text,
            "method": method_used,
            "cached_at": datetime.now().isoformat()
        }, f)

def count_stats(text):
    if not text:
        return 0, 0
    words = len(text.split())
    chars = len(text)
    return words, chars

def estimate_quality_score(markdown_text, method_used):
    """Estimate conversion quality (0-100)."""
    score = 50  # Base score
    
    # Method bonus
    if "Primary" in method_used:
        score += 30
    elif "cached" in method_used:
        score += 20
    
    # Content analysis
    if markdown_text:
        # Has tables
        if "|" in markdown_text and "---" in markdown_text:
            score += 10
        # Has headers
        if "#" in markdown_text:
            score += 5
        # Has formatting
        if "**" in markdown_text or "*" in markdown_text:
            score += 5
    
    return min(100, score)

def get_pdf_metadata(file_path):
    try:
        doc = fitz.open(file_path)
        metadata = doc.metadata
        page_count = len(doc)
        doc.close()
        
        info = []
        if metadata.get("title"):
            info.append(f"**Title:** {metadata['title']}")
        if metadata.get("author"):
            info.append(f"**Author:** {metadata['author']}")
        if metadata.get("subject"):
            info.append(f"**Subject:** {metadata['subject']}")
        info.append(f"**Pages:** {page_count}")
        if metadata.get("creationDate"):
            date_str = metadata["creationDate"]
            if date_str.startswith("D:"):
                date_str = date_str[2:10]
                try:
                    formatted = f"{date_str[:4]}-{date_str[4:6]}-{date_str[6:8]}"
                    info.append(f"**Created:** {formatted}")
                except:
                    pass
        
        return "\n".join(info) if info else "No metadata available"
    except Exception as e:
        return f"Could not read metadata: {e}"

def get_pdf_preview(file_path, page_num=0):
    """Generate preview image of PDF page."""
    try:
        doc = fitz.open(file_path)
        if page_num >= len(doc):
            page_num = 0
        page = doc[page_num]
        pix = page.get_pixmap(dpi=150)
        img_data = pix.tobytes("png")
        doc.close()
        
        # Save to temp file
        temp_path = os.path.join(tempfile.gettempdir(), f"preview_{page_num}.png")
        with open(temp_path, "wb") as f:
            f.write(img_data)
        return temp_path
    except Exception as e:
        print(f"Preview error: {e}")
        return None

def extract_images_from_pdf(file_path):
    """Extract all images from PDF."""
    images = []
    try:
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page = doc[page_num]
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                
                # Save image
                img_filename = f"page{page_num+1}_img{img_index+1}.{base_image['ext']}"
                img_path = os.path.join(IMAGES_DIR, img_filename)
                with open(img_path, "wb") as f:
                    f.write(image_bytes)
                images.append(img_path)
        
        doc.close()
    except Exception as e:
        print(f"Image extraction error: {e}")
    
    return images

def markdown_to_html(markdown_text):
    html = markdown2.markdown(markdown_text, extras=["tables", "fenced-code-blocks"])
    return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; 
                max-width: 800px; margin: 40px auto; padding: 20px; line-height: 1.6; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 12px; text-align: left; }}
        th {{ background-color: #f5f5f5; }}
        code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; }}
        pre {{ background: #f4f4f4; padding: 15px; border-radius: 5px; overflow-x: auto; }}
    </style>
</head>
<body>
{html}
</body>
</html>"""

def markdown_to_docx(markdown_text, output_path):
    doc = Document()
    for line in markdown_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)
    doc.save(output_path)

def markdown_to_txt(markdown_text):
    import re
    text = re.sub(r'#{1,6}\s', '', markdown_text)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text

def get_openrouter_cost(model_name):
    """Get cost per 1M tokens for OpenRouter model."""
    model_costs = {
        "Nemotron Nano 12B VL (FREE) ⭐": 0.0,
        "Gemini 2.0 Flash Lite ($0.08/1K pages)": 0.000075,
        "Qwen 2.5-VL 32B ($0.05/1K pages)": 0.00005,
        "Qwen 2.5-VL 72B ($0.15/1K pages)": 0.00015,
        "Mistral Pixtral Large ($2/1K pages)": 0.002
    }
    return model_costs.get(model_name, 0.0)

def estimate_cost(model_name, pages=1):
    """Estimate cost for processing given number of pages."""
    cost_per_1m = get_openrouter_cost(model_name)
    if cost_per_1m == 0.0:
        return "💰 **Estimated Cost:** FREE ⭐"
    
    # Estimate ~1000 tokens per page (conservative)
    tokens_estimate = pages * 1000
    cost_estimate = (tokens_estimate / 1_000_000) * cost_per_1m
    
    if cost_estimate < 0.01:
        return f"💰 **Estimated Cost:** ~${cost_estimate:.4f} ({pages} page{'s' if pages > 1 else ''})"
    else:
        return f"💰 **Estimated Cost:** ~${cost_estimate:.2f} ({pages} page{'s' if pages > 1 else ''})"

def toggle_settings(engine_choice):
    """Toggle visibility of OpenRouter vs RapidOCR settings."""
    is_openrouter = "OpenRouter" in engine_choice
    return gr.update(visible=is_openrouter), gr.update(visible=not is_openrouter)

def process_single_file(file_path, dpi, ocr_lang, page_start, page_end, use_cache, ocr_engine="RapidOCR", openrouter_model="free", openrouter_api_key=None, progress_callback=None):
    # Check cache
    if use_cache and file_path.lower().endswith(".pdf"):
        file_hash = get_file_hash(file_path)
        cache_key = f"{ocr_engine}_{openrouter_model if 'OpenRouter' in ocr_engine else dpi}"
        cached = get_cached_result(file_hash, cache_key, "md")
        if cached:
            return cached["markdown"], cached["method"] + " (cached)"
    
    markdown_text = None
    method_used = None
    
    if progress_callback:
        progress_callback(0.3, desc=f"🔍 Extracting {Path(file_path).name}...")
    
    # Route based on OCR engine selection
    if "OpenRouter" in ocr_engine:
        # Use OpenRouter for OCR
        if progress_callback:
            progress_callback(0.4, desc=f"🌐 Extracting with OpenRouter...")
        
        # Map UI model names to tier keys
        model_map = {
            "Nemotron Nano 12B VL (FREE) ⭐": "free",
            "Gemini 2.0 Flash Lite ($0.08/1K pages)": "cheap",
            "Qwen 2.5-VL 32B ($0.05/1K pages)": "balanced",
            "Qwen 2.5-VL 72B ($0.15/1K pages)": "quality",
            "Mistral Pixtral Large ($2/1K pages)": "premium"
        }
        model_tier = model_map.get(openrouter_model, "free")
        
        markdown_text, metadata = structure_engine.extract_with_openrouter(
            file_path, model=model_tier, api_key=openrouter_api_key
        )
        method_used = f"OpenRouter ({metadata.get('model_used', 'Unknown')})"
        
        # Check for errors
        if "Error" in markdown_text:
            # Fallback to RapidOCR if OpenRouter fails
            if progress_callback:
                progress_callback(0.5, desc="⚠️ OpenRouter failed, falling back to RapidOCR...")
            markdown_text = structure_engine.extract_with_rapidocr(file_path, dpi=dpi, lang=ocr_lang)
            method_used = "RapidOCR (Fallback)"
    else:
        # Use local processing (pymupdf4llm + RapidOCR fallback)
        # PRIMARY MODE: pymupdf4llm
        if file_path.lower().endswith(".pdf"):
            try:
                pages = None
                if page_start is not None or page_end is not None:
                    doc = fitz.open(file_path)
                    total_pages = len(doc)
                    doc.close()
                    start = (page_start or 1) - 1
                    end = min(page_end or total_pages, total_pages)
                    pages = list(range(start, end))
                
                markdown_text = pymupdf4llm.to_markdown(
                    file_path,
                    page_chunks=False,
                    write_images=False,
                    dpi=dpi,
                    force_text=True,
                    pages=pages,
                )
                
                if markdown_text and len(markdown_text.strip()) > MIN_TEXT_THRESHOLD:
                    method_used = "Primary (pymupdf4llm)"
                else:
                    markdown_text = None
                    
            except Exception as e:
                print(f"Primary extraction failed: {e}")
                markdown_text = None
        
        # FALLBACK MODE: RapidOCR
        if markdown_text is None:
            if progress_callback:
                progress_callback(0.5, desc=f"🔬 OCR scanning {Path(file_path).name}...")
            markdown_text = structure_engine.extract_with_rapidocr(file_path, dpi=dpi, lang=ocr_lang)
            method_used = "Fallback (RapidOCR)"
    
    # Cache result
    if use_cache and markdown_text and file_path.lower().endswith(".pdf"):
        file_hash = get_file_hash(file_path)
        cache_key = f"{ocr_engine}_{openrouter_model if 'OpenRouter' in ocr_engine else dpi}"
        save_to_cache(file_hash, cache_key, "md", markdown_text, method_used)
    
    return markdown_text, method_used


def process_upload(files, export_format, dpi, ocr_lang, page_start, page_end, use_cache, ocr_engine, openrouter_model, openrouter_api_key, progress=gr.Progress()):
    if not files:
        return None, None, None, "Upload a PDF or image to get started.", "", gr.update(visible=False), "", None, [], ""
    
    if not isinstance(files, list):
        files = [files]
    
    start_time = time.time()
    results = []
    metadata_info = ""
    preview_image = None
    extracted_images = []
    quality_info = ""
    
    try:
        # Get metadata and preview for first PDF
        first_file = files[0].name if hasattr(files[0], 'name') else files[0]
        if first_file.lower().endswith(".pdf"):
            metadata_info = get_pdf_metadata(first_file)
            preview_image = get_pdf_preview(first_file, 0)
            extracted_images = extract_images_from_pdf(first_file)
        
        for idx, file in enumerate(files):
            input_path = file.name if hasattr(file, 'name') else file
            progress(idx / len(files), desc=f"📄 Processing {idx+1}/{len(files)}...")
            
            markdown_text, method_used = process_single_file(
                input_path, dpi, ocr_lang, page_start, page_end, use_cache,
                ocr_engine, openrouter_model, openrouter_api_key,
                lambda p, desc: progress((idx + p) / len(files), desc=desc)
            )
            
            if markdown_text:
                results.append({
                    'filename': Path(input_path).stem,
                    'markdown': markdown_text,
                    'method': method_used
                })
        
        if not results:
            return None, None, None, "❌ Failed to extract content.", "", gr.update(visible=False), metadata_info, preview_image, extracted_images, quality_info
        
        progress(0.9, desc="✨ Finalizing...")
        
        elapsed_time = time.time() - start_time
        total_words = sum(count_stats(r['markdown'])[0] for r in results)
        total_chars = sum(count_stats(r['markdown'])[1] for r in results)
        
        # Update stats
        update_stats(total_words, elapsed_time)
        
        # Quality score
        quality_score = estimate_quality_score(results[0]['markdown'], results[0]['method'])
        quality_info = f"🎯 **Quality Score:** {quality_score}/100"
        
        if len(results) == 1:
            markdown_text = results[0]['markdown']
            method_used = results[0]['method']
            filename = results[0]['filename']
            
            temp_dir = tempfile.mkdtemp()
            
            if export_format == "Markdown (.md)":
                output_path = os.path.join(temp_dir, f"{filename}.md")
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(markdown_text)
            elif export_format == "HTML (.html)":
                output_path = os.path.join(temp_dir, f"{filename}.html")
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(markdown_to_html(markdown_text))
            elif export_format == "Plain Text (.txt)":
                output_path = os.path.join(temp_dir, f"{filename}.txt")
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(markdown_to_txt(markdown_text))
            elif export_format == "Word Document (.docx)":
                output_path = os.path.join(temp_dir, f"{filename}.docx")
                markdown_to_docx(markdown_text, output_path)
            
            stats_text = f"📊 **{total_words:,}** words • **{total_chars:,}** characters"
            status = f"✅ Converted using {method_used} in **{elapsed_time:.1f}s**"
            
            save_history({
                'filename': filename,
                'timestamp': datetime.now().isoformat(),
                'words': total_words,
                'method': method_used
            })
            
            return markdown_text, markdown_text, output_path, status, stats_text, gr.update(visible=True), metadata_info, preview_image, extracted_images, quality_info
        
        else:
            # Batch processing
            temp_dir = tempfile.mkdtemp()
            zip_path = os.path.join(temp_dir, "converted_documents.zip")
            
            with zipfile.ZipFile(zip_path, 'w') as zipf:
                for result in results:
                    filename = result['filename']
                    md = result['markdown']
                    
                    if export_format == "Markdown (.md)":
                        file_path = os.path.join(temp_dir, f"{filename}.md")
                        with open(file_path, "w", encoding="utf-8") as f:
                            f.write(md)
                    elif export_format == "HTML (.html)":
                        file_path = os.path.join(temp_dir, f"{filename}.html")
                        with open(file_path, "w", encoding="utf-8") as f:
                            f.write(markdown_to_html(md))
                    elif export_format == "Plain Text (.txt)":
                        file_path = os.path.join(temp_dir, f"{filename}.txt")
                        with open(file_path, "w", encoding="utf-8") as f:
                            f.write(markdown_to_txt(md))
                    elif export_format == "Word Document (.docx)":
                        file_path = os.path.join(temp_dir, f"{filename}.docx")
                        markdown_to_docx(md, file_path)
                    
                    zipf.write(file_path, os.path.basename(file_path))
            
            first_md = results[0]['markdown']
            stats_text = f"📊 **{len(results)} files** • **{total_words:,}** words"
            status = f"✅ Batch converted {len(results)} files in **{elapsed_time:.1f}s**"
            
            return first_md, first_md, zip_path, status, stats_text, gr.update(visible=True), metadata_info, preview_image, extracted_images, quality_info
        
    except Exception as e:
        traceback.print_exc()
        return None, None, None, f"❌ Error: {str(e)}", "", gr.update(visible=False), metadata_info, preview_image, extracted_images, quality_info

def get_history_display():
    history = load_history()
    if not history:
        return "No recent conversions"
    
    lines = ["### Recent Conversions\n"]
    for entry in history[:5]:
        timestamp = datetime.fromisoformat(entry['timestamp']).strftime("%m/%d %H:%M")
        lines.append(f"- **{entry['filename']}** ({entry['words']:,} words) - {timestamp}")
    
    return "\n".join(lines)

def get_stats_display():
    stats = load_stats()
    return f"""### 📈 Usage Stats
- **Total Conversions:** {stats['total_conversions']}
- **Total Words:** {stats['total_words']:,}
- **Avg Time:** {stats['avg_time']:.1f}s"""

def clear_cache():
    import shutil
    if os.path.exists(CACHE_DIR):
        shutil.rmtree(CACHE_DIR)
        os.makedirs(CACHE_DIR, exist_ok=True)
    return "✅ Cache cleared!"

# CSS
custom_css = """
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600&display=swap');
* { font-family: 'Inter', sans-serif !important; }
.container { max-width: 1200px !important; margin: auto; padding-top: 2rem; }
#header { text-align: center; margin-bottom: 1.5rem; }
#header h1 {
    font-weight: 600; font-size: 2.5rem; margin-bottom: 0.25rem;
    background: linear-gradient(135deg, #2563EB, #7C3AED);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent;
}
#header p { font-size: 1rem; opacity: 0.6; }
.primary-btn {
    background: linear-gradient(135deg, #2563EB, #1D4ED8) !important;
    border: none !important; color: white !important; font-weight: 600 !important;
    padding: 12px 24px !important; transition: all 0.2s ease;
}
.primary-btn:hover { transform: translateY(-2px); box-shadow: 0 8px 20px rgba(37, 99, 235, 0.4); }
.copy-btn { background: #10B981 !important; color: white !important; }
.theme-btn { background: #6B7280 !important; color: white !important; }
.input-card {
    border: 1px solid var(--border-color-primary); border-radius: 16px !important;
    padding: 1.5rem !important; background: var(--background-fill-secondary);
    box-shadow: 0 4px 12px rgba(0,0,0,0.05); transition: all 0.3s ease;
}
.input-card:has(input[type="file"]:hover) {
    border-color: #2563EB; box-shadow: 0 0 0 3px rgba(37, 99, 235, 0.1); transform: translateY(-2px);
}
.output-markdown { 
    padding: 1.5rem; border-radius: 12px; background: var(--background-fill-primary);
    border: 1px solid var(--border-color-primary); min-height: 350px;
}
.stats-box {
    padding: 0.75rem 1rem; background: linear-gradient(135deg, #F0F9FF, #E0F2FE);
    border-radius: 8px; border: 1px solid #BAE6FD; font-size: 0.9rem; margin-bottom: 1rem;
}
.metadata-box {
    padding: 0.75rem 1rem; background: linear-gradient(135deg, #FEF3C7, #FDE68A);
    border-radius: 8px; border: 1px solid #FCD34D; font-size: 0.85rem;
}
.quality-box {
    padding: 0.75rem 1rem; background: linear-gradient(135deg, #D1FAE5, #A7F3D0);
    border-radius: 8px; border: 1px solid #6EE7B7; font-size: 0.9rem;
}
.sidebar { background: var(--background-fill-secondary); border-radius: 12px; padding: 1rem; font-size: 0.85rem; }
.bottom-accent {
    height: 4px; width: 100%; background: linear-gradient(90deg, #2563EB, #7C3AED, #EC4899);
    position: fixed; bottom: 0; left: 0; z-index: 1000;
}
footer { visibility: hidden; }
"""

copy_js = """
function copyToClipboard() {
    const textArea = document.querySelector('textarea');
    if (textArea && textArea.value) {
        navigator.clipboard.writeText(textArea.value).then(() => {
            const btn = document.querySelector('.copy-btn');
            if (btn) {
                const originalText = btn.textContent;
                btn.textContent = '✓ Copied!';
                setTimeout(() => { btn.textContent = originalText; }, 2000);
            }
        });
    }
    return [];
}
"""

# Gradio Interface
with gr.Blocks(title="DocFlow", js=copy_js) as demo:
    
    gr.HTML('<div class="bottom-accent"></div>')

    with gr.Row(elem_classes="container"):
        
        with gr.Column(scale=5):
            
            with gr.Row():
                with gr.Column(elem_id="header", scale=3):
                    gr.Markdown("# DocFlow")
                    gr.Markdown("RAG-Optimized PDF & Image to Markdown Conversion")
                with gr.Column(scale=1):
                    theme_btn = gr.Button("🌙 Dark Mode", elem_classes="theme-btn", size="sm")
            
            # Settings
            with gr.Accordion("⚙️ Settings", open=False):
                # OCR Engine Selection
                with gr.Row():
                    ocr_engine = gr.Dropdown(
                        choices=["OpenRouter (Cloud - Recommended ⭐)", "RapidOCR (Local - Layout Aware)"],
                        value="OpenRouter (Cloud - Recommended ⭐)",
                        label="OCR Engine",
                        info="OpenRouter supports 100+ languages including Myanmar. RapidOCR features local layout analysis."
                    )
                
                # OpenRouter Settings (conditional)
                with gr.Group(visible=True) as openrouter_settings:
                    gr.Markdown("### 🌐 OpenRouter Settings")
                    with gr.Row():
                        openrouter_model = gr.Dropdown(
                            choices=[
                                "Nemotron Nano 12B VL (FREE) ⭐",
                                "Gemini 2.0 Flash Lite ($0.08/1K pages)",
                                "Qwen 2.5-VL 32B ($0.05/1K pages)",
                                "Qwen 2.5-VL 72B ($0.15/1K pages)",
                                "Mistral Pixtral Large ($2/1K pages)"
                            ],
                            value="Nemotron Nano 12B VL (FREE) ⭐",
                            label="Model",
                            info="FREE model recommended for most use cases"
                        )
                    with gr.Row():
                        openrouter_api_key = gr.Textbox(
                            label="OpenRouter API Key",
                            type="password",
                            placeholder="sk-or-v1-...",
                            info="Get free key at https://openrouter.ai"
                        )
                    cost_estimate = gr.Markdown("💰 **Estimated Cost:** FREE", elem_classes="cost-display")
                
                # RapidOCR Settings (conditional)
                with gr.Group(visible=False) as rapidocr_settings:
                    gr.Markdown("### 🔧 RapidOCR Settings")
                    with gr.Row():
                        dpi_slider = gr.Slider(minimum=150, maximum=600, value=300, step=50, label="DPI")
                        ocr_lang = gr.Dropdown(
                            choices=["en", "ch_sim", "ch_tra", "ja", "ko", "ru"],
                            value="en",
                            label="OCR Language",
                            info="Limited to 6 languages"
                        )
                
                # Common Settings
                gr.Markdown("### 📤 Export Options")
                with gr.Row():
                    export_format = gr.Dropdown(
                        choices=["Markdown (.md)", "HTML (.html)", "Plain Text (.txt)", "Word Document (.docx)"],
                        value="Markdown (.md)", label="Export Format"
                    )
                    use_cache = gr.Checkbox(value=True, label="Use Cache")
                with gr.Row():
                    page_start = gr.Number(label="Start Page", precision=0)
                    page_end = gr.Number(label="End Page", precision=0)

            
            # Info boxes
            with gr.Row():
                with gr.Column(scale=1):
                    metadata_display = gr.Markdown("", elem_classes="metadata-box", visible=False)
                with gr.Column(scale=1):
                    quality_display = gr.Markdown("", elem_classes="quality-box", visible=False)
            
            with gr.Row(elem_classes="input-card"):
                with gr.Column(scale=2):
                    file_input = gr.File(
                        label="📁 Upload Document(s)",
                        file_types=[".pdf", ".png", ".jpg", ".jpeg"],
                        type="filepath", file_count="multiple", height=120
                    )
                
                with gr.Column(scale=1):
                    process_btn = gr.Button("🚀 Convert", variant="primary", elem_classes="primary-btn")
                    status_box = gr.Markdown("Upload PDF(s) or image(s) to get started.")
            
            stats_display = gr.Markdown("", elem_classes="stats-box", visible=False)
            
            gr.Markdown("### Result")
            
            with gr.Tabs():
                with gr.TabItem("📄 Preview"):
                    output_md_view = gr.Markdown(elem_classes="output-markdown")
                with gr.TabItem("📝 Raw Code"):
                    output_raw_text = gr.TextArea(label="Markdown Source", lines=18)
                    copy_btn = gr.Button("📋 Copy to Clipboard", elem_classes="copy-btn", size="sm")
                with gr.TabItem("🖼️ Extracted Images"):
                    image_gallery = gr.Gallery(label="Images from PDF", columns=3, height=400)

            download_btn = gr.File(label="📥 Download", interactive=False)
        
        # Sidebar
        with gr.Column(scale=2):
            # PDF Preview
            pdf_preview = gr.Image(label="📄 PDF Preview", type="filepath", interactive=False)
            
            gr.Markdown("---")
            
            history_display = gr.Markdown(get_history_display())
            refresh_btn = gr.Button("🔄 Refresh History", size="sm")
            
            gr.Markdown("---")
            
            stats_panel = gr.Markdown(get_stats_display())
            
            gr.Markdown("---")
            
            clear_cache_btn = gr.Button("🗑️ Clear Cache", size="sm")
            cache_status = gr.Markdown("")

    # Events
    # Settings toggle
    ocr_engine.change(
        fn=toggle_settings,
        inputs=[ocr_engine],
        outputs=[openrouter_settings, rapidocr_settings]
    )
    
    # Dynamic cost estimation
    openrouter_model.change(
        fn=estimate_cost,
        inputs=[openrouter_model],
        outputs=[cost_estimate]
    )
    
    process_btn.click(
        fn=process_upload,
        inputs=[file_input, export_format, dpi_slider, ocr_lang, page_start, page_end, use_cache, 
                ocr_engine, openrouter_model, openrouter_api_key],
        outputs=[output_md_view, output_raw_text, download_btn, status_box, stats_display, stats_display, 
                 metadata_display, pdf_preview, image_gallery, quality_display]
    ).then(
        fn=get_history_display, outputs=[history_display]
    ).then(
        fn=get_stats_display, outputs=[stats_panel]
    ).then(
        fn=lambda m: gr.update(visible=bool(m)), inputs=[metadata_display], outputs=[metadata_display]
    ).then(
        fn=lambda q: gr.update(visible=bool(q)), inputs=[quality_display], outputs=[quality_display]
    )
    
    refresh_btn.click(fn=get_history_display, outputs=[history_display])
    clear_cache_btn.click(fn=clear_cache, outputs=[cache_status])
    copy_btn.click(fn=None, js="copyToClipboard")
    theme_btn.click(fn=None, js="() => { document.body.classList.toggle('dark'); return []; }")

if __name__ == "__main__":
    print("=" * 70)
    print("DocFlow - Production Ready PDF to Markdown Converter")
    print("=" * 70)
    print("Web UI: http://localhost:7860")
    print("API Server: Run 'python api.py' for REST API on port 8000")
    print("=" * 70)
    demo.launch(
        server_name="0.0.0.0", 
        server_port=7860,
        theme=gr.themes.Soft(primary_hue="blue", spacing_size="sm", radius_size="lg"),
        css=custom_css
    )
